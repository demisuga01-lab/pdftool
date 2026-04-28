import asyncio
import csv
import json
import logging
import shutil
import zipfile
from pathlib import Path
from typing import Any

from fastapi import HTTPException, status

logger = logging.getLogger(__name__)


class OCRService:
    async def _run_command(
        self,
        command: list[str],
        *,
        failure_status: int = status.HTTP_500_INTERNAL_SERVER_ERROR,
    ) -> tuple[str, str]:
        logger.info("Running command: %s", " ".join(command))
        try:
            process = await asyncio.create_subprocess_exec(
                *command,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await process.communicate()
        except FileNotFoundError as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Required command not found: {command[0]}",
            ) from exc
        except OSError as exc:
            raise HTTPException(
                status_code=failure_status,
                detail=f"Failed to run command: {command[0]}",
            ) from exc

        stdout_text = stdout.decode("utf-8", errors="replace")
        stderr_text = stderr.decode("utf-8", errors="replace")
        if process.returncode != 0:
            raise HTTPException(
                status_code=failure_status,
                detail=stderr_text.strip() or stdout_text.strip() or f"{command[0]} exited with code {process.returncode}",
            )
        return stdout_text, stderr_text

    def _ensure_parent(self, output_path: str | Path) -> Path:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        return output

    def _password_error_message(self) -> str:
        return "This PDF is password-protected. Please enter the password before running OCR."

    async def _pdf_is_encrypted(self, input_path: Path) -> bool:
        def read() -> bool:
            from pypdf import PdfReader

            return bool(PdfReader(str(input_path)).is_encrypted)

        try:
            return await asyncio.to_thread(read)
        except Exception:
            return False

    def _looks_like_password_error(self, text: str) -> bool:
        normalized = text.lower()
        return any(
            phrase in normalized
            for phrase in [
                "cannot authenticate password",
                "invalid password",
                "password required",
                "needs a password",
                "encrypted",
            ]
        )

    def _normalize_output_format(self, output_format: str) -> str:
        normalized = output_format.lower().strip()
        aliases = {"text": "txt", "searchable_pdf": "pdf", "searchable-pdf": "pdf"}
        normalized = aliases.get(normalized, normalized)
        if normalized not in {"txt", "json", "pdf", "docx", "hocr"}:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="output_format must be one of: txt, json, searchable_pdf, pdf, docx, hocr",
            )
        return normalized

    def _parse_tsv_content(self, tsv_content: str, page_number: int, text: str) -> dict[str, Any]:
        reader = csv.DictReader(tsv_content.splitlines(), delimiter="\t")
        words: list[dict[str, Any]] = []
        confidences: list[float] = []

        for row in reader:
            word_text = (row.get("text") or "").strip()
            if row.get("level") != "5" or not word_text:
                continue
            confidence = float(row.get("conf") or -1)
            if confidence >= 0:
                confidences.append(confidence)
            words.append(
                {
                    "text": word_text,
                    "confidence": confidence,
                    "left": int(row.get("left") or 0),
                    "top": int(row.get("top") or 0),
                    "width": int(row.get("width") or 0),
                    "height": int(row.get("height") or 0),
                }
            )

        return {
            "page": page_number,
            "text": text,
            "confidence": round(sum(confidences) / len(confidences), 2) if confidences else None,
            "words": words,
        }

    async def _render_pdf_pages(self, input_path: Path, output_dir: Path, dpi: int, password: str | None = None) -> list[Path]:
        safe_dpi = max(72, min(int(dpi or 300), 600))
        page_pattern = output_dir / "page-%04d.png"
        if await self._pdf_is_encrypted(input_path) and not password:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=self._password_error_message())

        render_source = input_path
        if password and await self._pdf_is_encrypted(input_path) and shutil.which("qpdf"):
            decrypted = output_dir / "decrypted-input.pdf"
            try:
                await self._run_command(["qpdf", f"--password={password}", "--decrypt", str(input_path), str(decrypted)])
                if decrypted.exists() and decrypted.stat().st_size > 0:
                    render_source = decrypted
            except HTTPException:
                render_source = input_path

        command = ["mutool", "draw"]
        if password and render_source == input_path:
            command.extend(["-p", password])
        command.extend(["-r", str(safe_dpi), "-F", "png", "-o", str(page_pattern), str(render_source)])

        try:
            process = await asyncio.create_subprocess_exec(
                *command,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await process.communicate()
        except FileNotFoundError as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Required command not found: mutool") from exc

        pages = sorted(output_dir.glob("page-*.png"))
        stderr_text = stderr.decode("utf-8", errors="replace").strip()
        stdout_text = stdout.decode("utf-8", errors="replace").strip()
        if process.returncode != 0 and not pages:
            detail = stderr_text or stdout_text or "Failed to render PDF pages before OCR"
            if self._looks_like_password_error(detail):
                detail = self._password_error_message()
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=detail)
        if not pages:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to render PDF pages before OCR",
            )
        return pages

    async def _ocr_text(self, image_path: Path, language: str) -> str:
        stdout, _ = await self._run_command(["tesseract", str(image_path), "stdout", "-l", language])
        return stdout.strip()

    async def _ocr_page_json(self, image_path: Path, page_number: int, language: str) -> dict[str, Any]:
        text = await self._ocr_text(image_path, language)
        output_base = image_path.with_name(f"ocr-page-{page_number:04d}")
        await self._run_command(["tesseract", str(image_path), str(output_base), "-l", language, "tsv"])
        return self._parse_tsv_content(output_base.with_suffix(".tsv").read_text(encoding="utf-8"), page_number, text)

    async def _write_docx(self, pages: list[dict[str, Any]], output_path: Path) -> str:
        output = self._ensure_parent(output_path)

        def write() -> None:
            from docx import Document

            document = Document()
            for page in pages:
                document.add_heading(f"Page {page['page']}", level=1)
                document.add_paragraph(str(page.get("text") or ""))
            document.save(output)

        await asyncio.to_thread(write)
        return str(output)

    async def _write_hocr(self, page_images: list[Path], output_dir: Path, language: str) -> str:
        hocr_paths: list[Path] = []
        for page_number, page_image in enumerate(page_images, start=1):
            output_base = output_dir / f"ocr-page-{page_number:04d}"
            await self._run_command(["tesseract", str(page_image), str(output_base), "-l", language, "hocr"])
            hocr_paths.append(output_base.with_suffix(".hocr"))

        if len(hocr_paths) == 1:
            return str(hocr_paths[0])

        archive_path = self._ensure_parent(output_dir / "ocr-hocr-pages.zip")

        def zip_hocr() -> None:
            with zipfile.ZipFile(archive_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
                for hocr_path in hocr_paths:
                    archive.write(hocr_path, arcname=hocr_path.name)

        await asyncio.to_thread(zip_hocr)
        return str(archive_path)

    async def _write_searchable_pdf(self, page_images: list[Path], output_dir: Path, language: str) -> str:
        page_pdfs: list[Path] = []
        for page_number, page_image in enumerate(page_images, start=1):
            output_base = output_dir / f"ocr-page-{page_number:04d}"
            await self._run_command(["tesseract", str(page_image), str(output_base), "-l", language, "pdf"])
            page_pdfs.append(output_base.with_suffix(".pdf"))

        if len(page_pdfs) == 1:
            output = self._ensure_parent(output_dir / "ocr-searchable.pdf")
            await asyncio.to_thread(page_pdfs[0].replace, output)
            return str(output)

        merged_output = self._ensure_parent(output_dir / "ocr-searchable.pdf")
        await self._run_command(["qpdf", "--empty", "--pages", *[str(path) for path in page_pdfs], "--", str(merged_output)])
        return str(merged_output)

    async def ocr(
        self,
        input_path: str | Path,
        output_dir: str | Path,
        language: str = "eng",
        output_format: str = "txt",
        dpi: int = 300,
        password: str | None = None,
    ) -> str:
        input_file = Path(input_path)
        output_directory = Path(output_dir)
        output_directory.mkdir(parents=True, exist_ok=True)
        normalized_output = self._normalize_output_format(output_format)

        if input_file.suffix.lower() == ".pdf":
            page_images = await self._render_pdf_pages(input_file, output_directory, dpi, password=password)
        else:
            page_images = [input_file]

        if normalized_output == "pdf":
            return await self._write_searchable_pdf(page_images, output_directory, language)

        if normalized_output == "hocr":
            return await self._write_hocr(page_images, output_directory, language)

        pages: list[dict[str, Any]] = []
        if normalized_output == "json":
            for page_number, page_image in enumerate(page_images, start=1):
                pages.append(await self._ocr_page_json(page_image, page_number, language))
        else:
            for page_number, page_image in enumerate(page_images, start=1):
                pages.append({"page": page_number, "text": await self._ocr_text(page_image, language)})

        if normalized_output == "docx":
            return await self._write_docx(pages, output_directory / "ocr-output.docx")

        combined = "\n\n".join(f"--- Page {page['page']} ---\n{page.get('text') or ''}" for page in pages).strip()
        if normalized_output == "json":
            output = self._ensure_parent(output_directory / "ocr-output.json")
            payload = {"pages": pages, "text": combined}
            await asyncio.to_thread(output.write_text, json.dumps(payload, indent=2), "utf-8")
            return str(output)

        output = self._ensure_parent(output_directory / "ocr-output.txt")
        await asyncio.to_thread(output.write_text, combined, "utf-8")
        return str(output)
